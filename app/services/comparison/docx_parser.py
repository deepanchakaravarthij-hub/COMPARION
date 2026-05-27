from __future__ import annotations

import hashlib
from dataclasses import dataclass
from io import BytesIO
from typing import Any

from docx import Document
from docx.enum.shape import WD_INLINE_SHAPE_TYPE  # type: ignore[import-untyped]


@dataclass(frozen=True)
class DocxStyle:
    paragraph_style: str | None = None
    font: str | None = None
    size_pt: float | None = None
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    color: str | None = None
    alignment: str | None = None
    numbering: str | None = None

    def signature(self) -> tuple[Any, ...]:
        return (
            self.paragraph_style,
            self.font,
            self.size_pt,
            self.bold,
            self.italic,
            self.underline,
            self.color,
            self.alignment,
            self.numbering,
        )


@dataclass(frozen=True)
class DocxRun:
    text: str
    style: DocxStyle
    source_ref: dict[str, Any]


@dataclass(frozen=True)
class DocxParagraph:
    text: str
    runs: list[DocxRun]
    style: DocxStyle
    source_ref: dict[str, Any]

    @property
    def normalized_text(self) -> str:
        return normalize_text(self.text)


@dataclass(frozen=True)
class DocxTableCell:
    text: str
    source_ref: dict[str, Any]

    @property
    def normalized_text(self) -> str:
        return normalize_text(self.text)


@dataclass(frozen=True)
class DocxTable:
    rows: list[list[DocxTableCell]]
    source_ref: dict[str, Any]


@dataclass(frozen=True)
class DocxImage:
    sha256: str
    content_type: str | None
    filename: str | None
    source_ref: dict[str, Any]
    blob: bytes
    page: int
    bbox: tuple[float, float, float, float] | None = None


@dataclass(frozen=True)
class DocxDocumentModel:
    paragraphs: list[DocxParagraph]
    tables: list[DocxTable]
    images: list[DocxImage]


def parse_docx(content: bytes, document_label: str) -> DocxDocumentModel:
    document = Document(BytesIO(content))
    paragraphs: list[DocxParagraph] = []
    tables: list[DocxTable] = []

    _append_paragraphs(paragraphs, document.paragraphs, document_label, "body")
    _append_tables(tables, document.tables, document_label, "body")

    for section_index, section in enumerate(document.sections, start=1):
        _append_paragraphs(
            paragraphs,
            section.header.paragraphs,
            document_label,
            f"header-{section_index}",
        )
        _append_tables(tables, section.header.tables, document_label, f"header-{section_index}")
        _append_paragraphs(
            paragraphs,
            section.footer.paragraphs,
            document_label,
            f"footer-{section_index}",
        )
        _append_tables(tables, section.footer.tables, document_label, f"footer-{section_index}")

    page_count = max(1, min(10, (len(paragraphs) // 25) + 1))
    package_images = _extract_images(document, document_label, page_count)
    inline_images = _extract_inline_images(document, document_label, page_count)
    return DocxDocumentModel(
        paragraphs=paragraphs,
        tables=tables,
        images=_merge_docx_images(package_images, inline_images),
    )


def normalize_text(text: str) -> str:
    return " ".join(text.split())


def _append_paragraphs(
    target: list[DocxParagraph],
    paragraphs: Any,
    document_label: str,
    part: str,
) -> None:
    for paragraph in paragraphs:
        paragraph_index = len(target) + 1
        source_ref = {
            "document": document_label,
            "page": None,
            "part": part,
            "block_id": f"{part}:p:{paragraph_index}",
            "paragraph": paragraph_index,
        }
        paragraph_style = _paragraph_style(paragraph)
        runs = [
            DocxRun(
                text=run.text,
                style=_run_style(run, paragraph_style),
                source_ref={**source_ref, "run": run_index},
            )
            for run_index, run in enumerate(paragraph.runs, start=1)
        ]
        target.append(
            DocxParagraph(
                text=paragraph.text,
                runs=runs,
                style=paragraph_style,
                source_ref=source_ref,
            )
        )


def _append_tables(
    target: list[DocxTable],
    tables: Any,
    document_label: str,
    part: str,
) -> None:
    for table in tables:
        table_index = len(target) + 1
        table_ref = {
            "document": document_label,
            "page": None,
            "part": part,
            "block_id": f"{part}:tbl:{table_index}",
            "table": table_index,
        }
        rows: list[list[DocxTableCell]] = []
        for row_index, row in enumerate(table.rows, start=1):
            cells = []
            for column_index, cell in enumerate(row.cells, start=1):
                source_ref = {
                    **table_ref,
                    "row": row_index,
                    "column": column_index,
                }
                cells.append(DocxTableCell(text=cell.text, source_ref=source_ref))
            rows.append(cells)
        target.append(DocxTable(rows=rows, source_ref=table_ref))


def _extract_images(document: Any, document_label: str, page_count: int) -> list[DocxImage]:
    images: list[DocxImage] = []
    relationships = sorted(document.part.rels.values(), key=lambda rel: rel.rId)
    for relationship in relationships:
        if "image" not in relationship.reltype:
            continue
        target_part = relationship.target_part
        blob = target_part.blob
        images.append(
            DocxImage(
                sha256=hashlib.sha256(blob).hexdigest(),
                content_type=getattr(target_part, "content_type", None),
                filename=getattr(relationship, "target_ref", None),
                source_ref={
                    "document": document_label,
                    "page": 1,
                    "part": "package",
                    "block_id": f"package:image:{len(images) + 1}",
                    "image": len(images) + 1,
                },
                blob=blob,
                page=1,
                bbox=(0.1, 0.1, 0.25, 0.25),
            )
        )
    return images


def _extract_inline_images(document: Any, document_label: str, page_count: int) -> list[DocxImage]:
    page_width_emu = 9_144_000
    page_height_emu = 6_858_000
    images: list[DocxImage] = []
    inline_shapes = list(document.inline_shapes)
    inline_count = max(len(inline_shapes), 1)
    for index, inline_shape in enumerate(inline_shapes, start=1):
        if inline_shape.type != WD_INLINE_SHAPE_TYPE.PICTURE:
            continue
        try:
            blip = inline_shape._inline.graphic.graphicData.pic.blipFill.blip
            embed = blip.embed
            part = document.part.related_parts[embed]
            blob = part.blob
        except AttributeError:
            continue
        page = max(1, min(page_count, (index * page_count) // inline_count))
        width = int(inline_shape.width) if inline_shape.width else page_width_emu // 4
        height = int(inline_shape.height) if inline_shape.height else page_height_emu // 6
        bbox = (
            0.1,
            min(0.85, 0.04 * (index - 1)),
            min(0.8, width / page_width_emu),
            min(0.5, height / page_height_emu),
        )
        images.append(
            DocxImage(
                sha256=hashlib.sha256(blob).hexdigest(),
                content_type=getattr(part, "content_type", None),
                filename=None,
                source_ref={
                    "document": document_label,
                    "page": page,
                    "part": "body",
                    "block_id": f"inline:image:{index}",
                    "image": index,
                },
                blob=blob,
                page=page,
                bbox=bbox,
            )
        )
    return images


def _merge_docx_images(
    package_images: list[DocxImage],
    inline_images: list[DocxImage],
) -> list[DocxImage]:
    merged: list[DocxImage] = []
    seen: set[str] = set()
    for image in [*inline_images, *package_images]:
        if image.sha256 in seen:
            continue
        seen.add(image.sha256)
        merged.append(image)
    return merged


def _paragraph_style(paragraph: Any) -> DocxStyle:
    return DocxStyle(
        paragraph_style=_style_name(paragraph),
        alignment=_enum_name(paragraph.alignment),
        numbering=_numbering(paragraph),
    )


def _run_style(run: Any, paragraph_style: DocxStyle) -> DocxStyle:
    font = run.font
    return DocxStyle(
        paragraph_style=paragraph_style.paragraph_style,
        font=font.name,
        size_pt=round(font.size.pt, 2) if font.size else None,
        bold=font.bold,
        italic=font.italic,
        underline=font.underline,
        color=str(font.color.rgb) if font.color and font.color.rgb else None,
        alignment=paragraph_style.alignment,
        numbering=paragraph_style.numbering,
    )


def _style_name(item: Any) -> str | None:
    style = getattr(item, "style", None)
    name = getattr(style, "name", None)
    return str(name) if name else None


def _enum_name(value: Any) -> str | None:
    if value is None:
        return None
    name = getattr(value, "name", None)
    return str(name) if name else str(value)


def _numbering(paragraph: Any) -> str | None:
    paragraph_properties = getattr(paragraph._p, "pPr", None)
    numbering_properties = getattr(paragraph_properties, "numPr", None)
    if numbering_properties is None:
        return None

    num_id = getattr(getattr(numbering_properties, "numId", None), "val", None)
    level = getattr(getattr(numbering_properties, "ilvl", None), "val", None)
    if num_id is None and level is None:
        return None
    return f"num:{num_id}:level:{level}"
